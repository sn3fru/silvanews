#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
================================================================================
|    PIPELINE DE ANÁLISE DE NOTÍCIAS - USABILIDADE INTELIGENTE (v6.3)      |
================================================================================
| MELHORIAS DESTA VERSÃO (v6.3 - Resumos Dinâmicos e Relatório Interativo):|
| 1. Sistema Hierárquico de Relevância:                                    |
|    - P1_CRITICO: Impacto imediato em decisões estratégicas e financeiras |
|    - P2_ESTRATEGICO: Tendências setoriais e oportunidades médio prazo    |
|    - P3_MONITORAMENTO: Contexto de mercado e informações gerais          |
|                                                                            |
| 2. Resumos Dinâmicos por Prioridade:                                     |
|    - P1_CRITICO: Resumos executivos detalhados (2-3 parágrafos)          |
|    - P2_ESTRATEGICO: Resumos padrão densos (1 parágrafo)                 |
|    - P3_MONITORAMENTO: Resumos concisos (1-2 frases)                     |
|                                                                            |
| 3. Relatório Interativo no Word:                                         |
|    - Títulos recolhíveis com navegação por seções                        |
|    - Painel de navegação ativo para acesso rápido                        |
|    - Guia de leitura integrado no documento                              |
|    - Experiência de usuário otimizada para consumo executivo             |
|                                                                            |
| 4. Performance e Robustez Mantidas:                                      |
|    - Processamento paralelo preservado (10+15 workers)                   |
|    - Mapeamento robusto por IDs e cache inteligente                      |
|    - Formatação profissional de fontes únicas                           |
--------------------------------------------------------------------------------
"""


import os
import re
import json
import glob
import time
from datetime import datetime
from typing import List, Dict, Any, Optional, Literal

# Importações para Manipulação de Arquivos
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Importação da API do Google Gemini
from google import genai
from google.genai import types

# Importações do Pydantic para Validação de Dados
from pydantic import BaseModel, Field, ValidationError

# Importação do PyMuPDF para Chunking de PDFs
import fitz  # PyMuPDF

# Importação para processamento paralelo
import concurrent.futures

# --- CONSTANTES DE CONFIGURAÇÃO ---
PDF_DIRECTORY = 'pdfs'
OUTPUT_DIRECTORY = 'relatorios_gerados'
CACHE_DIRECTORY = 'cache_noticias'

# --- CONFIGURAÇÕES ANTI-TRUNCAMENTO ---
PAGINAS_POR_CHUNK = 5  # Número de páginas por pedaço ao dividir PDFs grandes
LIMITE_PAGINAS_CHUNKING = 50  # PDFs com mais páginas serão automaticamente divididos

# Importação para arquivos temporários únicos
import tempfile
import threading

# ==============================================================================
# 1. FUNÇÕES AUXILIARES PARA MIGRAÇÃO DE CACHE
# ==============================================================================

def migrar_noticia_cache_legado(noticia_data):
    """
    Migra dados de cache antigos para o novo formato com relevance_score e relevance_reason.
    Se os campos estiverem ausentes, adiciona valores padrão baseados na prioridade.
    """
    # Se já tem os campos novos, retorna como está
    if 'relevance_score' in noticia_data and 'relevance_reason' in noticia_data:
        return noticia_data
    
    # Migração baseada na prioridade
    prioridade = noticia_data.get('prioridade', 'P3_MONITORAMENTO')
    
    # Adiciona valores padrão baseados na prioridade
    if prioridade == 'P1_CRITICO':
        noticia_data['relevance_score'] = 85.0
        noticia_data['relevance_reason'] = "Migrado do cache: Classificado como P1_CRITICO"
    elif prioridade == 'P2_ESTRATEGICO':
        noticia_data['relevance_score'] = 65.0  
        noticia_data['relevance_reason'] = "Migrado do cache: Classificado como P2_ESTRATEGICO"
    else:  # P3_MONITORAMENTO
        noticia_data['relevance_score'] = 35.0
        noticia_data['relevance_reason'] = "Migrado do cache: Classificado como P3_MONITORAMENTO"
    
    return noticia_data

def validar_e_migrar_cache(noticias_cache, nome_arquivo):
    """
    Valida e migra dados de cache, tentando recuperar dados antigos quando possível.
    """
    noticias_migradas = []
    
    for noticia_data in noticias_cache:
        try:
            # Tenta migrar dados antigos
            noticia_migrada = migrar_noticia_cache_legado(noticia_data)
            
            # Valida com o modelo Pydantic
            noticia_validada = Noticia(**noticia_migrada)
            noticias_migradas.append(noticia_validada.model_dump())
            
        except ValidationError as e:
            # Se ainda assim falhou na validação, pula esta notícia
            print(f"   ⚠️ Notícia inválida ignorada no cache de {nome_arquivo}: {e}")
            continue
        except Exception as e:
            print(f"   ⚠️ Erro ao migrar notícia do cache de {nome_arquivo}: {e}")
            continue
    
    return noticias_migradas

# ==============================================================================
# 2. MODELOS DE VALIDAÇÃO DE DADOS COM PYDANTIC
# ==============================================================================

class Noticia(BaseModel):
    """Modelo de validação para notícias extraídas dos PDFs."""
    titulo: str = Field(..., min_length=1, description="Título da notícia")
    texto_completo: str = Field(..., min_length=1, description="Texto completo da notícia")
    jornal: str = Field(..., min_length=1, description="Nome do jornal")
    autor: Optional[str] = Field(default="N/A", description="Autor da notícia")
    pagina: Optional[str] = Field(default=None, description="Página onde a notícia foi encontrada")
    data: Optional[str] = Field(default=None, description="Data de publicação")
    categoria: Optional[str] = Field(default=None, description="Categoria da notícia")
    # Tag de classificação (rigorosa - apenas as 4 válidas)
    tag: Literal['Governo e Politica', 'Economia e Tecnologia', 'Judicionario', 'Empresas Privadas']
    # NOVO CAMPO: Adiciona a prioridade para filtragem
    prioridade: Literal['P1_CRITICO', 'P2_ESTRATEGICO', 'P3_MONITORAMENTO'] = Field(..., description="Nível de prioridade da notícia")
    
    # --- NOVO CAMPO PARA A URL ---
    url: Optional[str] = Field(default=None, description="URL da fonte original da notícia")
    
    # --- CAMPOS PARA RANKING E EXPLAINABILITY (AGORA OPCIONAIS) ---
    relevance_score: Optional[float] = Field(default=None, ge=0, le=100, description="Score de 0 a 100 da relevância para a mesa de Special Situations.")
    relevance_reason: Optional[str] = Field(default=None, description="Justificativa curta em qual regra/assunto a notícia se encaixa.")
    
    def model_post_init(self, __context):
        """Adiciona valores padrão para campos ausentes após a validação inicial."""
        if self.relevance_score is None or self.relevance_reason is None:
            # Aplica migração baseada na prioridade
            if self.prioridade == 'P1_CRITICO':
                self.relevance_score = self.relevance_score or 85.0
                self.relevance_reason = self.relevance_reason or "Migrado: Classificado como P1_CRITICO"
            elif self.prioridade == 'P2_ESTRATEGICO':
                self.relevance_score = self.relevance_score or 65.0
                self.relevance_reason = self.relevance_reason or "Migrado: Classificado como P2_ESTRATEGICO"
            else:  # P3_MONITORAMENTO
                self.relevance_score = self.relevance_score or 35.0
                self.relevance_reason = self.relevance_reason or "Migrado: Classificado como P3_MONITORAMENTO"

class NoticiaResumida(BaseModel):
    """Modelo para notícias resumidas usadas no agrupamento."""
    titulo: str = Field(..., min_length=1, description="Título da notícia")
    jornal: str = Field(..., min_length=1, description="Nome do jornal")



class FonteResumo(BaseModel):
    """Modelo para fontes de um resumo."""
    jornal: Optional[str] = Field(default=None, description="Nome do jornal")
    pagina: Optional[str] = Field(default=None, description="Página da notícia")
    autor: Optional[str] = Field(default=None, description="Autor da notícia")
    # --- NOVO CAMPO PARA A URL DA FONTE ---
    url: Optional[str] = Field(default=None, description="URL da fonte")

class ResumoFinal(BaseModel):
    """Modelo de validação para resumos finais."""
    titulo_final: str = Field(..., min_length=1, description="Título final do resumo")
    resumo_final: str = Field(..., min_length=1, description="Resumo consolidado")
    fontes: Optional[List[FonteResumo]] = Field(default=[], description="Lista de fontes")
    tag: Literal['Governo e Politica', 'Economia e Tecnologia', 'Judicionario', 'Empresas Privadas']
    # NOVO CAMPO: Adiciona a prioridade para ordenação do relatório
    prioridade: Literal['P1_CRITICO', 'P2_ESTRATEGICO', 'P3_MONITORAMENTO'] = Field(..., description="Prioridade máxima do grupo de notícias")

# ==============================================================================
# 3. FUNÇÃO AUXILIAR PARA HIPERLINKS NO WORD
# ==============================================================================

def add_hyperlink(paragraph, text, url):
    """
    Adiciona um hiperlink a um parágrafo.

    Args:
        paragraph: O objeto de parágrafo do python-docx.
        text: O texto a ser exibido para o link.
        url: A URL para a qual o link aponta.

    Returns:
        O objeto 'run' do hiperlink.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estilo do link (azul e sublinhado)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

# ==============================================================================
# 4. LISTA HIERÁRQUICA DE RELEVÂNCIA (v6.0 - com Prioridades)
# ==============================================================================

LISTA_RELEVANCIA_HIERARQUICA = {
    "P1_CRITICO": {
        "descricao": "OPORTUNIDADES ACIONÁVEIS AGORA: Situações de estresse financeiro, M&A e arbitragem legal que demandam ação imediata.",
        "assuntos": [
            # Mapeado de f-2, f-3: O núcleo de Special Situations
            "Recuperação Judicial (RJ)", "Falência", "Pedido de Falência", "Assembleia de Credores",
            # Mapeado de f-4, f-5: Ativos estressados
            "Créditos Inadimplentes (NPLs)", "Créditos Podres (Distressed Debt)", "Venda de Carteira de NPL",
            # Mapeado de g-2, g-4: Oportunidades de arbitragem legal/fiscal
            "Crédito Tributário (teses, oportunidades de monetização)", "Disputas Societárias Relevantes",
            # Mapeado de f-6, f-7: Ativos específicos do governo
            "FCVS (Fundo de Compensação de Variações Salariais - apenas notícias sobre liquidação ou venda de créditos)",
            "Dívida Ativa (apenas notícias sobre venda de grandes blocos ou securitização)",
            # Mapeado de f-1 e feedback de classificados: Oportunidades imobiliárias de grande porte
            "Leilões Judiciais de Ativos (apenas imóveis ou participações societárias acima de R$10 milhões)",
            # Eventos corporativos críticos
            "Fusões e Aquisições (M&A) - Anunciadas ou em negociação avançada",
            "Crise de Liquidez Aguda", "Quebra de Covenants", "Default de Dívida"
        ],
        "empresas": [ # Empresas cujo estresse financeiro ou M&A é P1
             # Mapeado de j, k e feedback
            "Americanas S.A.", "Oi S.A.", "Casas Bahia (Grupo Pão de Açúcar)", "Light S.A.",
            "Gol Linhas Aéreas", "Azul Linhas Aéreas", "Petrobras", "Vale",
            "IRB Resseguradora (IRB Brasil RE)"
        ]
    },
    "P2_ESTRATEGICO": {
        "descricao": "MONITORAMENTO ESTRATÉGICO: Tendências e eventos que podem gerar oportunidades P1 no futuro.",
        "assuntos": [
            # Mapeado de g-1, g-2, g-3: Mudanças que podem gerar futuras teses
            "Mudanças em Legislação (Tributária, Societária, Falimentar, Precatórios)",
            # Mapeado de a-2, c-1, c-2, c-3, b-1: Setores de alto capital e tecnologia
            "Inteligência Artificial (IA - apenas grandes movimentos de mercado, M&A no setor ou regulação pesada)",
            "Semicondutores (geopolítica da cadeia de suprimentos, grandes investimentos/fábricas)",
            "Energia Nuclear (grandes projetos, concessões, marco regulatório)",
            "Aeroespacial e Defesa (grandes contratos governamentais, privatizações)",
            # Política e regulação com impacto direto
            "Política Econômica (Decisões de juros e política fiscal que afetem o crédito e a saúde financeira das empresas)",
            "Decisões do CADE (bloqueio de fusões, imposição de remédios)",
            "Ativismo Acionário (grandes investidores tentando influenciar a gestão)"
        ],
        "empresas": [ # Empresas para monitoramento contínuo de resultados e movimentos estratégicos
            # Mapeado de h (Big Techs)
            "Alphabet", "AMD", "Apple", "Google", "Intel", "Intuitive Machines", "Meta",
            "Micron Technology", "Microsoft", "Netflix", "Tesla", "Nvidia",
            # Mapeado de i (Energia)
            "Constellation Energy Group", "Siemens Energy AG",
            # Mapeado de j, k (Bancos e Seguros)
            "Banco Master", "Banco Pan", "Caixa Econômica Federal", "PREVIC"
        ]
    },
    "P3_MONITORAMENTO": {
        "descricao": "CONTEXTO DE MERCADO: Informações gerais para entendimento do cenário macro, sem ação direta.",
        "assuntos": [
            # Mapeado de d (Cripto)
            "Criptomoedas (apenas visão macro de mercado, adoção institucional ou regulação. Sem análise técnica de moedas específicas).",
            # Cenário internacional
            "Geoeconomia", "Acordos Comerciais (Mercosul-UE, etc.)", "Decisões do FED e BCE",
            # Mapeado de e-1: Games (apenas M&A no setor)
            "Games (apenas notícias sobre grandes fusões e aquisições, ex: Microsoft comprando Activision)"
        ],
        "empresas": [],
        "regras_especiais": {
            "balancos_financeiros": "Notícias sobre balanços só são relevantes se (A) forem de empresas listadas em P2, ou (B) indicarem um estresse financeiro severo (risco de RJ/Falência), tornando-as P1.",
            "classificados_e_leiloes": "Anúncios de classificados são 99% irrelevantes. A única exceção é um leilão judicial ou venda de um ativo único de altíssimo valor (>R$10M), que deve ser tratado como P1."
        }
    }
}

# Criar uma versão em string para ser usada nos prompts
def gerar_lista_relevancia_para_prompt():
    texto_prompt = ""
    for prioridade, data in LISTA_RELEVANCIA_HIERARQUICA.items():
        texto_prompt += f"**{prioridade} ({data['descricao']})**\n"
        if data['assuntos']:
            texto_prompt += "- Assuntos: " + ", ".join(data['assuntos']) + "\n"
        if data['empresas']:
            texto_prompt += "- Empresas: " + ", ".join(data['empresas']) + "\n"
        texto_prompt += "\n"
    return texto_prompt

LISTA_RELEVANCIA_FORMATADA = gerar_lista_relevancia_para_prompt()

# ==============================================================================
# 5. PROMPTS DETALHADOS PARA O PIPELINE DE IA (v5.0)
# ==============================================================================

# # Prompt original mais restritivo para uso em caso de necessidade
# PROMPT_EXTRACAO_TOLERANCIA_ZERO_V7 = f"""
# Sua identidade: Você é um analista sênior da mesa de 'Special Situations' do banco BTG Pactual. Seu bônus depende da qualidade do seu filtro.

# **A REGRA DE OURO (SEU FILTRO MENTAL):**
# Ao ler uma notícia, faça a si mesmo esta única pergunta: **"Isto aponta para um ativo estressado, uma assimetria de informação ou uma ineficiência de mercado que pode ser monetizada HOJE?"** Se a resposta não for um "SIM" óbvio, **REJEITE IMEDIATAMENTE.**

# --------------------------------------------------------------------------------
# **ETAPA 1: ANÁLISE E CLASSIFICAÇÃO**
# --------------------------------------------------------------------------------
# 1.  Avalie a notícia contra a **LISTA DE REJEIÇÃO IMEDIATA**. Se houver correspondência, descarte.
# 2.  Avalie contra a **LISTA DE INTERESSES**. Se não houver correspondência clara, descarte.
# 3.  Se a notícia passar, você DEVE atribuir um **Score de Relevância** e uma **Justificativa**.

# **LISTA DE INTERESSES (FOCO EXCLUSIVO):**
# {LISTA_RELEVANCIA_FORMATADA}

# **LISTA DE REJEIÇÃO IMEDIATA (SE A NOTÍCIA FOR SOBRE ISTO, É LIXO):**
# - **Fofoca Política e Partidária:** Disputas no congresso, popularidade de políticos, eleições, vida pessoal de figuras públicas.
# - **Legislação Social e Ambiental Genérica:** Leis de cotas, regras de EAD, Lei Rouanet, ESG, créditos de carbono, licenciamento ambiental. (Exceção: se levar diretamente à falência de uma empresa relevante).
# - **Operações do Dia a Dia do Governo:** Restituição de IR, filas do INSS, dados do IBGE, protocolos de blitz.
# - **Crimes, Golpes e Segurança Pública:** Golpes virtuais, fraudes, disputas criminais de baixo impacto.
# - **Eventos, Seminários, Cultura e Esportes:** Cobertura de fóruns, notícias sobre games (exceto M&A), parques temáticos, crises em clubes de futebol.
# - **Anúncios Publicitários e Classificados Genéricos.**

# --------------------------------------------------------------------------------
# **ETAPA 2: EXTRAÇÃO (APENAS SE APROVADO)**
# --------------------------------------------------------------------------------
# Se a notícia for relevante, extraia os dados no formato JSON abaixo.

# - **`relevance_score`**: Score de 0-100. P1 (RJ, Falência, M&A direto) deve ser > 85. P2 (Regulatório, Tendências) entre 50-85. P3 (Contexto Macro) < 50.
# - **`relevance_reason`**: Justifique sua decisão em uma frase. Ex: "Encaixa-se na regra P1-CRITICO: Recuperação Judicial".
# - **`tag`**: Use uma das 4 tags: 'Governo e Politica', 'Economia e Tecnologia', 'Judicionario', 'Empresas Privadas'.

# **FORMATO DE SAÍDA (JSON PURO):**
# ```json
# [
#   {{
#     "titulo": "...",
#     "texto_completo": "Resumo focado na tese de investimento...",
#     "jornal": "...",
#     "autor": "...",
#     "pagina": "...",
#     "data": "...",
#     "categoria": "O item mais específico da LISTA DE INTERESSES (ex: Recuperação Judicial)",
#     "tag": "Uma das 4 tags válidas",
#     "prioridade": "A prioridade correta (P1_CRITICO, P2_ESTRATEGICO, ou P3_MONITORAMENTO)",
#     "relevance_score": 95.0,
#     "relevance_reason": "Encaixa-se na regra P1-CRITICO: Pedido de Falência de empresa relevante do setor aéreo."
#   }}
# ]
# ```
# Seja brutal no filtro. Se nenhuma notícia passar, retorne [].
# """

# Prompt mais permissivo para extração inicial (melhor captura)
PROMPT_EXTRACAO_PERMISSIVO_V8 = f"""
Sua identidade: Você é um analista junior da mesa de 'Special Situations' do banco BTG Pactual. Sua função é fazer uma primeira triagem ampla de notícias.

**OBJETIVO:** Capturar TODAS as notícias que possam ter alguma relevância para decisões de investimento, mesmo que remotamente. 
É melhor incluir uma notícia desnecessária do que perder uma oportunidade importante.

--------------------------------------------------------------------------------
**LISTA DE REJEIÇÃO IMEDIATA (SE FOR SOBRE ISSO, É RUÍDO):**
--------------------------------------------------------------------------------
- **Crimes e Segurança Pública Cotidiana:** Prisões individuais (como a do homem em Roraima), estatísticas genéricas de crimes (feminicídios, latrocínios do Anuário de Segurança), operações policiais de rotina, fraudes comuns. **Exceção:** Apenas se envolver diretamente uma empresa P1 ou P2 em um esquema de corrupção de grande escala.
- **Esportes:** Finanças de clubes (como a falta de verba do Corinthians), resultados de jogos, contratações. **Exceção:** Apenas se houver um processo de Recuperação Judicial ou M&A de uma SAF (Sociedade Anônima do Futebol) relevante.
- **Cultura, Fofoca e Entretenimento:** Vidas de celebridades (morte de Hulk Hogan), festivais (Dança de Joinville), lançamentos de musicais, moda (Flavia Aranha), disputas judiciais de natureza pessoal (Juliana Oliveira vs. SBT).
- **Assuntos Locais sem Impacto Sistêmico:** Disputas de bairro (prédio em Higienópolis), classificados e leilões de baixo valor (Rico Leilão de veículos).
- **Política Partidária Pura:** Disputas internas de partidos, fofocas de bastidores, popularidade de políticos. **Exceção:** Decisões de política econômica com impacto direto.

**FOCO PRINCIPAL - CAPTURE SE A NOTÍCIA FOR SOBRE:**
{LISTA_RELEVANCIA_FORMATADA}

**REJEITE APENAS SE FOR CLARAMENTE IRRELEVANTE:**
- Esportes (exceto aspectos financeiros de clubes)
- Entretenimento e cultura (exceto aspectos de negócio)
- Crimes cotidianos sem impacto empresarial
- Política partidária pura (exceto políticas econômicas)
- Eventos sociais e celebridades

**INSTRUÇÕES DE EXTRAÇÃO:**
1. **Seja GENEROSO** na classificação inicial - prefira incluir do que excluir
2. Para notícias de fronteira, classifique como P3_MONITORAMENTO 
3. Use scores mais baixos (30-50) para notícias incertas, mas INCLUA elas
4. O filtro rigoroso será aplicado later no pipeline

**INSTRUÇÕES DE EXTRAÇÃO E SCORING:**
- **`prioridade`**: Atribua P1, P2 ou P3 RIGOROSAMENTE com base na lista. Uma notícia sobre "Recuperação Judicial" é SEMPRE `P1_CRITICO`.
- **`relevance_score`**: Seja conservador.
    - **P1_CRITICO (Score 85-100):** Apenas para eventos acionáveis (RJ, Falência, M&A anunciado).
    - **P2_ESTRATEGICO (Score 50-84):** Apenas para tendências setoriais e mudanças regulatórias claras.
    - **P3_MONITORAMENTO (Score 20-49):** Para contexto macroeconômico e notícias de empresas monitoradas que não se encaixam em P1/P2.
- **`tag`**: Classifique com precisão


**FORMATO DE SAÍDA (JSON PURO):**
```json
[
  {{
    "titulo": "...",
    "texto_completo": "Resumo da notícia...",
    "jornal": "...",
    "autor": "...",
    "pagina": "...",
    "data": "...",
    "categoria": "Categoria da LISTA DE INTERESSES ou 'Geral'",
    "tag": "Uma das 4 tags: 'Governo e Politica', 'Economia e Tecnologia', 'Judicionario', 'Empresas Privadas'",
    "prioridade": "P1_CRITICO, P2_ESTRATEGICO, ou P3_MONITORAMENTO",
    "relevance_score": 45.0,
    "relevance_reason": "Notícia sobre setor relevante para acompanhamento"
  }}
]
```

**LEMBRE-SE:** É melhor capturar 100 notícias e depois filtrar para 20, do que capturar apenas 5 e perder oportunidades importantes.
"""

# Prompt para os resumos de 1 parágrafo do TOP 12
PROMPT_RESUMO_CRITICO_V1 = """
Você é um analista de investimentos escrevendo um briefing para o comitê executivo.
Sua tarefa é resumir o seguinte evento em um **único parágrafo conciso de, no máximo, 5 linhas.**
Foque nos fatos essenciais: Quem, O quê, Quando, Onde e Qual a implicação financeira ou oportunidade de negócio.
Ignore detalhes secundários. Seja direto e informativo.

DADOS DO EVENTO PARA RESUMIR:
{DADOS_DO_GRUPO}

FORMATO DE SAÍDA OBRIGATÓRIO (JSON PURO):
```json
{{
  "resumo_final": "Seu resumo executivo de um parágrafo aqui."
}}
```
"""

# Prompt para os resumos de 1 linha do RADAR
PROMPT_RADAR_MONITORAMENTO_V1 = """
Você está criando um "Radar de Monitoramento" para executivos. Sua tarefa é transformar as notícias do cluster abaixo em UM ÚNICO bullet point de UMA LINHA, começando com a entidade principal.

IMPORTANTE: Crie APENAS um bullet point por cluster, consolidando todas as notícias relacionadas em uma única linha informativa.

Exemplos:
- Cluster sobre iFood: "iFood: em negociações para adquirir a Alelo por R$ 5 bilhões"
- Cluster sobre governo: "Mercado de Carbono: governo adia a criação de agência reguladora para o setor"

DADOS DO CLUSTER PARA TRANSFORMAR EM BULLET POINT:
{DADOS_DO_GRUPO}

FORMATO DE SAÍDA OBRIGATÓRIO (JSON PURO):
```json
{{
  "bullet_point": "Entidade: resumo consolidado das notícias do cluster em uma linha"
}}
```
"""


PROMPT_AGRUPAMENTO_CONSOLIDADO_V2 = """
Você é um especialista em análise de conteúdo focado em granularidade. Sua tarefa é agrupar notícias de uma lista JSON que se referem exatamente ao mesmo fato gerador. Diferentes jornais cobrirão o mesmo fato com títulos distintos, e sua missão é identificá-los com precisão.

**DIRETRIZES DE AGRUPAMENTO:**

1.  **INTEGRIDADE TOTAL:** TODAS as notícias da entrada DEVEM ser alocadas a um grupo. Notícias sem par formam um grupo de um único item. NENHUMA notícia pode ser descartada.

2.  **FOCO NO NÚCLEO SEMÂNTICO:** O que realmente aconteceu? Qual foi a decisão, o anúncio ou o evento?
    - **EXEMPLO DE AGRUPAR (MESMO FATO):**
      - Notícia 1: "Alexandre de Moraes decide não prender Jair Bolsonaro, mas confirma descumprimento de restrições."
      - Notícia 2: "Ministro do STF não decreta prisão preventiva de Jair Bolsonaro, mas mantém e esclarece cautelares."
      - **Análise:** O núcleo semântico é o mesmo: a decisão de Moraes de não prender Bolsonaro, mantendo as restrições. **Devem estar no mesmo grupo.**

3.  **DESDOBRAMENTOS PODEM SER AGRUPADOS NA MESMA NOTICIA DEIXANDO:** Uma ação e a reação a ela são O MESMO FATO EM MOMENTOS DIFERENTES.
      - Grupo A: "Governo anuncia nova política de preços."
      - Grupo B: "Setor industrial critica nova política de preços."

4.  **TEMA PRINCIPAL PRECISO:** O `tema_principal` deve descrever o fato gerador de forma neutra e específica. Deve ser o título da ação consolidada.

5.  **MAPEAMENTO POR ID:** Use os `ids_originais` fornecidos para garantir a correspondência.

**FORMATO DE ENTRADA (EXEMPLO):**
[
 {"id": 0, "titulo": "Apple lança iPhone 20", "jornal": "Jornal Tech", "trecho": "..."},
 {"id": 1, "titulo": "Novo iPhone 20 da Apple chega ao mercado", "jornal": "Jornal Varejo", "trecho": "..."},
 {"id": 2, "titulo": "Analistas reagem ao lançamento do iPhone 20", "jornal": "Jornal Mercado", "trecho": "..."},
 {"id": 3, "titulo": "Tesla anuncia novo carro elétrico", "jornal": "Jornal Auto", "trecho": "..."}
]

**FORMATO DE SAÍDA OBRIGATÓRIO (JSON PURO - USANDO 'ids_originais'):**
```json
[
 {
   "tema_principal": "Apple lança o novo iPhone 20",
   "ids_originais": [0, 1]
 },
 {
   "tema_principal": "Alexandre de Moraes manda prender Jair Bolsonaro",
   "ids_originais": [2]
 },
 {
   "tema_principal": "Tesla anuncia novo modelo de carro elétrico",
   "ids_originais": [3]
 }
]
```
"""

PROMPT_RESUMO_FINAL_V3 = """
Você é um analista de inteligência criando um resumo sobre um evento específico, baseado em um CLUSTER de notícias relacionadas. A profundidade do seu resumo deve variar conforme o **Nível de Detalhe** solicitado.

**IMPORTANTE:** Você está resumindo um CLUSTER DE NOTÍCIAS sobre o mesmo fato gerador, não sessões separadas. Combine todas as informações das notícias do cluster em um resumo coerente e abrangente.

**NÍVEIS DE DETALHE:**
-   **Executivo (P1_CRITICO):** Um resumo aprofundado de 2 a 3 parágrafos. Detalhe o contexto, os principais dados (valores, percentuais), os players envolvidos e as implicações estratégicas (riscos/oportunidades). Seja completo.
-   **Padrão (P2_ESTRATEGICO):** Um único parágrafo denso e informativo que sintetiza os fatos mais importantes do evento, combinando informações de todas as notícias do cluster.
-   **Conciso (P3_MONITORAMENTO):** Uma única frase (máximo duas) que captura a essência do evento consolidado.

**MISSÃO:**
Baseado no CLUSTER de notícias fornecido e no **Nível de Detalhe** `{NIVEL_DE_DETALHE}` solicitado, produza um resumo consolidado que integre todas as informações do cluster.

**FORMATO DE SAÍDA OBRIGATÓRIO (JSON PURO):**
```json
{{
  "titulo_final": "Use exatamente o tema_principal fornecido no cluster.",
  "resumo_final": "O resumo consolidado de todas as notícias do cluster conforme o Nível de Detalhe especificado."
}}
```

**DADOS DO CLUSTER PARA ANÁLISE:**
{DADOS_DO_GRUPO}
"""

PROMPT_EXTRACAO_JSON_V1 = f"""
Você é um classificador e sumarizador de notícias. Sua função é analisar UM ÚNICO item de notícia (título e texto completo) que já foi pré-extraído de um site.

--------------------------------------------------------------------------------
**ETAPA 1: ANÁLISE DE RELEVÂNCIA E SUMARIZAÇÃO**
--------------------------------------------------------------------------------
1.  **Relevância:** Compare o conteúdo da notícia com a `LISTA DE RELEVÂNCIA` abaixo. A notícia DEVE ser sobre um dos tópicos de interesse e NENHUM dos tópicos proibidos.
2.  **Sumarização:** Se a notícia for relevante, crie um resumo informativo e bem estruturado do texto original, com no máximo 5 parágrafos, focando em fatos, dados e consequências. Se o texto original já for curto e objetivo, pode transcrevê-lo.

**LISTA DE RELEVÂNCIA (ALLOW LIST):**
{LISTA_RELEVANCIA_FORMATADA}

**LISTA DE ASSUNTOS PROIBIDOS (REJEIÇÃO IMEDIATA):**
- Cultura geral (exceto games), esportes, saúde local, política partidária genérica, crimes comuns.
- Entretenimento puro sem conexão tecnológica ou comercial
- Crimes comuns, acidentes locais, fofocas políticas sem impacto econômico
- Notícias puramente sociais ou culturais

--------------------------------------------------------------------------------
**ETAPA 2: AÇÃO**
--------------------------------------------------------------------------------
-   **SE a notícia for IRRELEVANTE:** Retorne um JSON com uma lista vazia: `[]`.
-   **SE a notícia for RELEVANTE:** Extraia as informações e as retorne como um JSON contendo uma lista com UM ÚNICO objeto, conforme o formato abaixo.

--------------------------------------------------------------------------------
**ETAPA 3: AUTO-VERIFICAÇÃO DA TAG (OBRIGATÓRIO)**
--------------------------------------------------------------------------------
Sua tarefa mais crítica é atribuir a tag correta. A `tag` DEVE ser UMA e APENAS UMA das 5 opções a seguir. Analise os exemplos positivos e negativos para cada uma antes de decidir.

**1. TAG: 'Internacional'**
- **Definição:** Eventos que ocorrem fora do Brasil ou envolvem a relação do Brasil com atores externos (países, empresas estrangeiras, bancos centrais globais).
- **Exemplos de ONDE USAR (Positivos):**
    - `França: Macron anuncia reconhecimento oficial do Estado da Palestina...`
    - `Estados Unidos e Israel: abandonam negociações de cessar-fogo em Gaza.`
    - `Donald Trump e Jerome Powell: confrontam-se publicamente na sede do Federal Reserve...`
    - `EUA: Fecham acordo comercial com Japão e avançam em negociações com a União Europeia...`
    - `LVMH: registra queda de 3% na receita... impactada por menor demanda na China e Japão.`
    - `Brasil: exportações de soja para a China seguem em alta...`
    - `Christine Lagarde (BCE): adota tom duro, levando mercado a rever apostas de cortes de juros...`
    - `Coreia do Sul: expande 'soft power' no Brasil via gastronomia...`
    - `Antonov An-24: queda na Rússia expõe os desafios da aviação russa com sanções...`
- **Exemplos de ONDE NÃO USAR (Negativos):**
    - Notícia: `IBGE: Acesso à internet e uso de celular no Brasil atingiram novos recordes...` -> **Tag Correta: 'Economia e Politica'**. (É sobre dados internos do Brasil).
    - Notícia: `Trigo: plantio no Sul do Brasil chega ao fim com projeções de queda...` -> **Tag Correta: 'Economia e Politica'**. (É sobre agricultura nacional).


**2. TAG: 'Economia e Politica' (Foco Brasil)**
- **Definição:** Acontecimentos da política e economia DOMÉSTICA do Brasil. Envolve governo federal, ministérios, políticas públicas e dados macroeconômicos nacionais.
- **Exemplos de ONDE USAR (Positivos):**
    - `IBGE: divulga dados do IPCA-15 de julho, importante indicador da inflação oficial do país.`
    - `Governo Brasileiro: finaliza plano de contingência e negocia com EUA para evitar tarifaço...`
    - `Arrecadação Federal: atinge recorde de R$ 234,59 bilhões...`
    - `Salários: estudo da LCA 4intelligence aponta que são o principal fator da inflação de serviços...`
- **Exemplos de ONDE NÃO USAR (Negativos):**
    - Notícia: `Alexandre de Moraes: decide não prender Jair Bolsonaro...` -> **Tag Correta: 'Legislativo e Judiciario'**. (Decisão do poder judiciário).
    - Notícia: `Donald Trump e Jerome Powell: confrontam-se...` -> **Tag Correta: 'Internacional'**. (Política e economia dos EUA).
    - Notícia: `Anuário Brasileiro de Segurança Pública: revela queda das mortes violentas...` -> **Rejeitar Notícia**. (Tema indesejado de segurança pública).


**3. TAG: 'Legislativo e Judiciario' (Foco Brasil)**
- **Definição:** Decisões, julgamentos e processos dos poderes Judiciário e Legislativo do Brasil. Envolve STF, STJ, TJ, processos de falência, recuperação judicial e votações no Congresso.
- **Exemplos de ONDE USAR (Positivos):**
    - `Alexandre de Moraes: decide não prender Jair Bolsonaro, mas confirma descumprimento de restrições...`
    - `Jair Bolsonaro: condenado pelo TJDF a pagar R$ 150 mil por danos morais...`
    - `STJ: decide que pagamento de legado de renda vitalícia não depende da conclusão do inventário.`
    - `Partido Verde (PV): aciona STF para contestar artigo da Lei Anticorrupção...`
    - `STF: suspende, por pedido de vista do ministro Flávio Dino, o julgamento...`
    - `W3 Camisetas Ltda.: falência decretada pela 2ª Vara Empresarial de Belo Horizonte/MG.`
    - `Belo Monte: obteve vitória judicial contra o ONS...`
- **Exemplos de ONDE NÃO USAR (Negativos):**
    - Notícia: `Governo Brasileiro: finaliza plano de contingência...` -> **Tag Correta: 'Economia e Politica'**. (Ação do poder Executivo).
    - Notícia: `Homem preso em Roraima ao sacar dinheiro roubado...` -> **Rejeitar Notícia**. (Crime comum, não é uma decisão judicial de relevância sistêmica).

**4. TAG: 'Tecnologia'**
- **Definição:** Notícias sobre inovação, inteligência artificial, semicondutores, data centers, cibersegurança e o modelo de negócios de empresas de tecnologia.
- **Exemplos de ONDE USAR (Positivos):**
    - `Conselho Nacional de Educação: elabora o primeiro regramento para o uso de Inteligência Artificial...`
    - `TikTok: investirá R$ 50 bilhões no Brasil, com Ministro Silveira prevendo mais aportes em tecnologia e IA...`
    - `Demanda por IA: eleva o custo de energia a recorde histórico...`
    - `Nvidia: US$ 1 bilhão em chips contrabandeados para a China...`
    - `Ações meme: impulsionam 'gamma squeezes' no mercado de opções...` (Fenômeno de mercado com base tecnológica).
- **Exemplos de ONDE NÃO USAR (Negativos):**
    - Notícia: `IBGE: Acesso à internet e uso de celular no Brasil...` -> **Tag Correta: 'Economia e Politica'**. (É um dado socioeconômico, não sobre a tecnologia em si).
    - Notícia: `São Paulo: concentra 18,5% dos roubos e furtos de celular do Brasil...` -> **Rejeitar Notícia**. (É sobre crime, não tecnologia).


**5. TAG: 'Empresas Privadas'**
- **Definição:** Ações e movimentos específicos de empresas, como investimentos, emissão de dívida, fusões e aquisições (M&A), disputas corporativas e projetos.
- **Exemplos de ONDE USAR (Positivos):**
    - `Empresas: Sabesp, Randon e Localiza estão entre as 26 companhias com ofertas de debêntures...`
    - `Cade: arquiva inquérito que investigava Globo, Disney e Warner...`
    - `Grupo Kalunga: assina acordo de sublicenciamento com a ESPN...`
    - `Helbor: Grupo de moradores de Higienópolis tenta barrar construção de prédio residencial...`
- **Exemplos de ONDE NÃO USAR (Negativos):**
    - Notícia: `LVMH: registra queda de 3% na receita... impactada por menor demanda na China e Japão.` -> **Tag Correta: 'Internacional'**. (A notícia é sobre o impacto de fatores geopolíticos/macroeconômicos na empresa).
    - Notícia: `W3 Camisetas Ltda.: falência decretada...` -> **Tag Correta: 'Legislativo e Judiciario'**. (O fato gerador é a decisão judicial, não uma ação da empresa).
**ATRIBUIÇÃO DE PRIORIDADE:** Ao extrair, você DEVE analisar a LISTA HIERÁRQUICA e atribuir a prioridade correta ('P1_CRITICO', 'P2_ESTRATEGICO', ou 'P3_MONITORAMENTO') ao campo "prioridade". Notícias que se encaixam em P1 devem ser sempre P1, mesmo que também toquem em temas P2 ou P3.

**FORMATO DE SAÍDA (JSON PURO, LISTA COM UM ÚNICO ITEM SE RELEVANTE):**
```json
[
  {{
    "titulo": "Título da notícia, conforme recebido",
    "autor": "N/A",
    "texto_completo": "O resumo completo e estruturado da notícia, com no máximo 5 parágrafos.",
    "pagina": "N/A",
    "jornal": "Será preenchido posteriormente",
    "data": "Será preenchido posteriormente",
    "categoria": "Categoria específica da LISTA DE RELEVÂNCIA (ex: Bitcoin, Apple, etc.)",
    "tag": "Uma das 4 tags válidas",
    "prioridade": "A prioridade correta (P1_CRITICO, P2_ESTRATEGICO, ou P3_MONITORAMENTO) baseada na LISTA HIERÁRQUICA"
  }}
]
```
INSTRUÇÃO FINAL: Analise o item de notícia fornecido. Se for relevante, retorne a lista com o resumo JSON. Se não, retorne [].
"""

PROMPT_CORRECAO_JSON = """
A seguinte string deveria ser um JSON válido contendo uma lista de objetos, mas contém um erro de sintaxe que impede sua decodificação.
Sua única tarefa é corrigir a sintaxe e retornar APENAS o código JSON válido e completo. Não adicione nenhum comentário ou texto explicativo.

JSON Quebrado:
{json_quebrado}
"""

# ==============================================================================
# 6. FUNÇÕES AUXILIARES E DE PROCESSAMENTO
# ==============================================================================

def corrigir_tag_invalida(tag_original: str) -> str:
    """
    Mapeia tags inválidas ou similares para uma das 4 tags válidas.
    
    Args:
        tag_original: Tag original que pode estar inválida
        
    Returns:
        Tag válida mapeada
    """
    if not tag_original or not isinstance(tag_original, str):
        return 'Empresas Privadas'  # Tag padrão
    
    tag_limpa = tag_original.strip().lower()
    
    # Mapeamento de tags similares para as corretas
    MAPEAMENTO_TAGS = {
        # Governo e Política
        'governo e politica': 'Governo e Politica',
        'governo e política': 'Governo e Politica', 
        'política': 'Governo e Politica',
        'politica': 'Governo e Politica',
        'governo': 'Governo e Politica',
        'política econômica': 'Governo e Politica',
        'politica economica': 'Governo e Politica',
        'política pública': 'Governo e Politica',
        'política publica': 'Governo e Politica',
        
        # Economia e Tecnologia  
        'economia e tecnologia': 'Economia e Tecnologia',
        'economia': 'Economia e Tecnologia',
        'tecnologia': 'Economia e Tecnologia',
        'tech': 'Economia e Tecnologia',
        'ia': 'Economia e Tecnologia',
        'inteligência artificial': 'Economia e Tecnologia',
        'inteligencia artificial': 'Economia e Tecnologia',
        'cripto': 'Economia e Tecnologia',
        'criptomoedas': 'Economia e Tecnologia',
        'bitcoin': 'Economia e Tecnologia',
        
        # Judiciário
        'judicionario': 'Judicionario',
        'judiciário': 'Judicionario',
        'judicial': 'Judicionario',
        'justiça': 'Judicionario',
        'justica': 'Judicionario',
        'tribunal': 'Judicionario',
        'stf': 'Judicionario',
        'stj': 'Judicionario',
        'falência': 'Judicionario',
        'falencia': 'Judicionario',
        'recuperação judicial': 'Judicionario',
        'recuperacao judicial': 'Judicionario',
        
        # Empresas Privadas
        'empresas privadas': 'Empresas Privadas',
        'empresa': 'Empresas Privadas',
        'empresas': 'Empresas Privadas',
        'corporativo': 'Empresas Privadas',
        'negócios': 'Empresas Privadas',
        'negocios': 'Empresas Privadas',
        'setor privado': 'Empresas Privadas',
        'm&a': 'Empresas Privadas',
        'fusões e aquisições': 'Empresas Privadas',
        'fusoes e aquisicoes': 'Empresas Privadas',
    }
    
    # Tenta mapeamento direto
    if tag_limpa in MAPEAMENTO_TAGS:
        return MAPEAMENTO_TAGS[tag_limpa]
    
    # Tenta mapeamento por palavras-chave
    if any(palavra in tag_limpa for palavra in ['governo', 'político', 'politico', 'ministério', 'ministerio']):
        return 'Governo e Politica'
    elif any(palavra in tag_limpa for palavra in ['tecnologia', 'economia', 'mercado', 'financeiro', 'cripto']):
        return 'Economia e Tecnologia'
    elif any(palavra in tag_limpa for palavra in ['judicial', 'tribunal', 'justiça', 'justica', 'falência', 'falencia']):
        return 'Judicionario'
    elif any(palavra in tag_limpa for palavra in ['empresa', 'corporativo', 'negócio', 'negocio', 'setor']):
        return 'Empresas Privadas'
    
    # Se nada funcionou, retorna tag padrão
    print(f"   ⚠️ Tag desconhecida '{tag_original}' mapeada para 'Empresas Privadas'")
    return 'Empresas Privadas'

def verificar_dependencias():
    """Verifica se todas as dependências necessárias estão instaladas."""
    dependencias_faltantes = []
    
    try:
        import google.generativeai
    except ImportError:
        dependencias_faltantes.append("google-generativeai")
    
    try:
        from docx import Document
    except ImportError:
        dependencias_faltantes.append("python-docx")
    
    try:
        from pydantic import BaseModel
    except ImportError:
        dependencias_faltantes.append("pydantic")
    
    try:
        import fitz  # PyMuPDF
    except ImportError:
        dependencias_faltantes.append("PyMuPDF")
    
    if dependencias_faltantes:
        print("❌ ERRO: Dependências não encontradas!")
        print(f"   Instale com: pip install {' '.join(dependencias_faltantes)}")
        return False
    
    return True

def contar_paginas_pdf(caminho_pdf: str) -> int:
    """
    Conta o número de páginas de um PDF.
    
    Args:
        caminho_pdf: Caminho para o arquivo PDF
        
    Returns:
        Número de páginas do PDF
    """
    try:
        with fitz.open(caminho_pdf) as doc:
            return doc.page_count
    except Exception as e:
        print(f"⚠️ Erro ao contar páginas de {caminho_pdf}: {e}")
        return 0

def limpar_nome_arquivo(nome: str) -> str:
    """Limpa uma string para ser usada como um nome de arquivo seguro."""
    nome = re.sub(r'[<>:"/\\|?*]', '_', nome)
    nome = re.sub(r'\s+', '_', nome)
    nome = re.sub(r'[^\w\-_.]', '', nome)
    return nome[:150]

def extrair_json_da_resposta(resposta: str) -> Any:
    """
    Tenta extrair e decodificar um objeto JSON de uma string de resposta do LLM,
    que pode estar envolto em markdown, texto solto ou ser truncado.
    Inclui depuração detalhada em caso de falha.
    """
    # ETAPA 0: Validação inicial da resposta
    # Garante que a resposta não é vazia ou nula antes de qualquer processamento.
    if not isinstance(resposta, str) or not resposta.strip():
        print("❌ Erro ao extrair JSON: A resposta recebida da API está vazia ou não é uma string.")
        return None

    json_str = ""
    # ETAPA 1: Tenta encontrar um bloco de código JSON explícito (```json ... ```)
    # Esta é a forma mais confiável de extrair o conteúdo.
    match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', resposta, re.DOTALL)
    if match:
        json_str = match.group(1).strip()
    else:
        # ETAPA 2 (Fallback): Se não houver bloco de código, busca o primeiro '[' ou '{'
        # Isso ajuda a ignorar qualquer texto introdutório que a API possa ter enviado.
        start_bracket = resposta.find('[')
        start_brace = resposta.find('{')
        
        start = -1
        if start_bracket != -1 and (start_bracket < start_brace or start_brace == -1):
            start = start_bracket
        elif start_brace != -1:
            start = start_brace

        if start != -1:
            # Se encontrou um marcador, assume que o resto da string é a tentativa de JSON.
            json_str = resposta[start:].strip()
        else:
            # SE NENHUM MARCADOR FOI ENCONTRADO, a resposta não se parece com JSON.
            print("❌ Erro ao extrair JSON: Nenhum marcador de início ('[' ou '{') foi encontrado.")
            print("📋 RESPOSTA COMPLETA DA API (para depuração):")
            print("-" * 50)
            print(resposta)
            print("-" * 50)
            return None

    # ETAPA 3: Tenta decodificar o JSON extraído
    # É aqui que erros de truncamento (JSON incompleto) serão pegos.
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        # Ocorreu um erro de decodificação. O JSON está malformado ou incompleto.
        print(f"❌ Erro ao decodificar JSON (malformado ou incompleto): {e}")
        print("📋 CONTEÚDO QUE FALHOU NA ANÁLISE (para depuração):")
        print("-" * 50)
        print(json_str)
        print("-" * 50)
        return None

def processar_pdf_com_chunking(caminho_pdf: str, client, generation_config_decision) -> List[Dict[str, Any]]:
    """
    Processa um PDF dividindo-o em chunks menores para evitar truncamento.
    Usa arquivos temporários únicos para evitar conflitos em processamento paralelo.
    """
    nome_arquivo = os.path.basename(caminho_pdf)
    print(f"📖 Processando PDF com chunking: `{nome_arquivo}`")
    noticias_consolidadas = []

    try:
        doc = fitz.open(caminho_pdf)
        total_paginas = doc.page_count
        print(f"   📄 Total de páginas: {total_paginas}")

        for i_chunk in range(0, total_paginas, PAGINAS_POR_CHUNK):
            start_page = i_chunk
            end_page = min(i_chunk + PAGINAS_POR_CHUNK, total_paginas)
            
            print(f"   ⚡ Processando chunk: páginas {start_page + 1} a {end_page}")
            
            # Criar um arquivo temporário único para este chunk
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_chunk_path = temp_file.name
            
            try:
                # Criar um PDF temporário para o chunk atual
                chunk_doc = fitz.open()
                chunk_doc.insert_pdf(doc, from_page=start_page, to_page=end_page - 1)
                chunk_doc.save(temp_chunk_path)
                chunk_doc.close()
                
                # Upload e processamento do chunk
                uploaded_file = client.files.upload(file=temp_chunk_path)
                
                # Aguarda processamento
                while uploaded_file.state.name == "PROCESSING":
                    time.sleep(5)
                    uploaded_file = client.files.get(name=uploaded_file.name)
                
                if uploaded_file.state.name != "ACTIVE":
                    print(f"     ❌ Erro no processamento do chunk: {uploaded_file.state.name}")
                    client.files.delete(name=uploaded_file.name)
                    continue
                
                response = client.models.generate_content(
                    model='gemini-2.0-flash-lite',
                    contents=[uploaded_file, PROMPT_EXTRACAO_PERMISSIVO_V8],
                    config=generation_config_decision
                )
                client.files.delete(name=uploaded_file.name)
                
                # ======================================================================
                # BLOCO DE CORREÇÃO (Tratamento de Resposta da API)
                # ======================================================================
                # O objetivo deste bloco é garantir que a resposta da API não foi bloqueada
                # ou retornou vazia antes de tentar processar o JSON.
                response_text = None
                try:
                    response_text = response.text
                except ValueError:
                    # Este erro é levantado pela biblioteca do Gemini quando o conteúdo é bloqueado.
                    print(f"     ⚠️ A resposta da API para o chunk foi bloqueada (provavelmente por filtros de segurança).")
                    if response.prompt_feedback and response.prompt_feedback.block_reason:
                         print(f"        -> Motivo do Bloqueio: {response.prompt_feedback.block_reason.name}")

                # Só prossegue se response_text for uma string válida
                if response_text:
                    noticias_chunk = extrair_json_da_resposta(response_text)
                    if noticias_chunk and isinstance(noticias_chunk, list):
                        noticias_validadas_chunk = []
                        for j, noticia_data in enumerate(noticias_chunk):
                            try:
                                if noticia_data.get('jornal'):
                                    noticia_data['jornal'] = nome_arquivo
                                
                                # CORREÇÃO DE TAG ANTES DA VALIDAÇÃO
                                if 'tag' in noticia_data:
                                    noticia_data['tag'] = corrigir_tag_invalida(noticia_data['tag'])
                                
                                # APLICA MIGRAÇÃO PARA GARANTIR CAMPOS NECESSÁRIOS
                                noticia_data = migrar_noticia_cache_legado(noticia_data)
                                
                                noticia_obj = Noticia(**noticia_data)
                                noticias_validadas_chunk.append(noticia_obj.model_dump())
                            except ValidationError as e:
                                print(f"     ⚠️ Notícia {j+1} do chunk inválida e DESCARTADA: {[error['loc'][0] for error in e.errors()]}")
                            except Exception as e:
                                print(f"     ⚠️ Erro ao validar notícia {j+1} do chunk: {e}")
                        
                        noticias_consolidadas.extend(noticias_validadas_chunk)
                        print(f"     ✅ {len(noticias_validadas_chunk)} notícias válidas extraídas do chunk")
                    else:
                        print(f"     ✅ 0 notícias relevantes encontradas no chunk (conforme esperado).")
                else:
                    print(f"     INFO: A API não retornou conteúdo para este chunk.")
                # ======================================================================
                # FIM DO BLOCO DE CORREÇÃO
                # ======================================================================
                        
            except Exception as e:
                print(f"     ❌ Erro ao processar chunk {start_page + 1}-{end_page}: {e}")
            finally:
                # Limpa arquivo temporário único
                try:
                    if os.path.exists(temp_chunk_path):
                        os.remove(temp_chunk_path)
                except Exception as cleanup_error:
                    print(f"     ⚠️ Aviso: Não foi possível limpar arquivo temporário: {cleanup_error}")
        
        doc.close()
        print(f"   🎯 Total consolidado: {len(noticias_consolidadas)} notícias de `{nome_arquivo}`")
        return noticias_consolidadas
        
    except Exception as e:
        print(f"❌ ERRO CRÍTICO no chunking de `{nome_arquivo}`: {e}")
        return []

def processar_uma_noticia(args: tuple) -> Optional[Dict[str, Any]]:
    """
    Função 'worker' para processar um único item de notícia.
    Projetada para ser usada em um processo paralelo.
    Inclui lógica de auto-correção de JSON.
    """
    noticia_original, client, generation_config_decision, nome_arquivo_origem = args
    i, noticia_data = noticia_original
    
    titulo = noticia_data.get("titulo")
    texto_completo = noticia_data.get("texto_completo")
    
    if not all([titulo, texto_completo]):
        print(f"   ⚠️ Notícia {i} pulada: 'titulo' ou 'texto_completo' ausente.")
        return None

    print(f"   -> Processando notícia {i}: '{titulo[:60]}...'")

    try:
        # Prepara o conteúdo para enviar à API
        conteudo_para_api = json.dumps({"titulo": titulo, "texto_completo": texto_completo}, ensure_ascii=False)
        
        # 1ª Tentativa: Chamada principal à API
        response = client.models.generate_content(
            model='gemini-2.0-flash-lite',
            contents=[PROMPT_EXTRACAO_JSON_V1, conteudo_para_api],
            config=generation_config_decision
        )
        
        response_text = response.text
        noticias_extraidas = extrair_json_da_resposta(response_text)

    except json.JSONDecodeError as e:
        print(f"   🐛 JSON malformado detectado na notícia {i}. Tentando auto-correção... Erro: {e}")
        try:
            # 2ª Tentativa: Chamada para o prompt de correção
            response_corrigida = client.models.generate_content(
                model='gemini-2.0-flash-lite',
                contents=[PROMPT_CORRECAO_JSON.format(json_quebrado=response_text)],
                config=generation_config_decision
            )
            noticias_extraidas = extrair_json_da_resposta(response_corrigida.text)
            print(f"   ✅ Auto-correção bem-sucedida para a notícia {i}.")
        except Exception as e_corr:
            print(f"   ❌ Falha na auto-correção para a notícia {i}. Descartando. Erro: {e_corr}")
            return None
    except Exception as e_main:
        print(f"   ❌ Erro inesperado ao processar a notícia {i}: {e_main}")
        return None

    # Validação e enriquecimento do resultado
    if noticias_extraidas and isinstance(noticias_extraidas, list) and len(noticias_extraidas) > 0:
        try:
            noticia_processada = noticias_extraidas[0]
            
            # CORREÇÃO: Tratar fonte corretamente para crawlers JSON
            if nome_arquivo_origem.endswith('.json'):
                # Para arquivos JSON, usar a fonte real do dicionário
                fonte_real = noticia_data.get('fonte', 'N/A')
                noticia_processada['jornal'] = fonte_real
                # Preservar o arquivo de origem para referência se necessário
                noticia_processada['arquivo_origem'] = nome_arquivo_origem
            else:
                # Para PDFs, usar o nome do arquivo
                noticia_processada['jornal'] = nome_arquivo_origem
            
            noticia_processada['data'] = noticia_data.get('data_publicacao')
            noticia_processada['autor'] = noticia_data.get('autor', 'N/A')
            
            # +++ INÍCIO DA ALTERAÇÃO +++
            # CAPTURAR A URL DO OBJETO ORIGINAL DA NOTÍCIA (DO ARQUIVO JSON)
            # Tenta capturar de 'url' ou 'link', que são os nomes mais comuns.
            url_original = noticia_data.get('url') or noticia_data.get('link')
            noticia_processada['url'] = url_original
            # +++ FIM DA ALTERAÇÃO +++
            
            # CORREÇÃO DE TAG ANTES DA VALIDAÇÃO
            if 'tag' in noticia_processada:
                noticia_processada['tag'] = corrigir_tag_invalida(noticia_processada['tag'])
            
            # APLICA MIGRAÇÃO PARA GARANTIR CAMPOS NECESSÁRIOS
            noticia_processada = migrar_noticia_cache_legado(noticia_processada)
            
            # Valida com Pydantic para garantir a estrutura final
            noticia_obj = Noticia(**noticia_processada)
            print(f"   ✅ Notícia {i} validada com sucesso.")
            return noticia_obj.model_dump()
        except ValidationError as e_val:
            print(f"   ⚠️ Notícia {i} DESCARTADA (erro de validação Pydantic): {[error['loc'][0] for error in e_val.errors()]}")
            return None
    
    return None

def processar_arquivo_json(caminho_json: str, client, generation_config_decision) -> List[Dict[str, Any]]:
    """
    Processa um arquivo JSON contendo uma lista de notícias de forma PARALELA e ROBUSTA.
    """
    import concurrent.futures

    nome_arquivo = os.path.basename(caminho_json)
    print(f"📖 Processando arquivo JSON com performance otimizada: `{nome_arquivo}`")
    
    try:
        with open(caminho_json, 'r', encoding='utf-8') as f:
            dados_json = json.load(f)
    except (json.JSONDecodeError, FileNotFoundError) as e:
        print(f"❌ Erro crítico ao ler o arquivo JSON `{nome_arquivo}`: {e}")
        return []

    if not isinstance(dados_json, list):
        print(f"❌ Erro: O arquivo JSON `{nome_arquivo}` não contém uma lista na raiz.")
        return []

    # Número de chamadas paralelas. Um bom ponto de partida é entre 10 e 20.
    # Não aumente demais para não sobrecarregar a API (rate limits).
    MAX_WORKERS = 15
    noticias_consolidadas = []
    
    # Prepara os argumentos para cada worker
    tarefas = [( (i, noticia), client, generation_config_decision, nome_arquivo) for i, noticia in enumerate(dados_json, 1)]

    print(f" 🚀 Iniciando processamento paralelo de {len(dados_json)} notícias com até {MAX_WORKERS} workers...")
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        # map executa as tarefas em paralelo e retorna os resultados na ordem
        resultados = executor.map(processar_uma_noticia, tarefas)
        
        # Coleta os resultados que não são nulos
        for resultado in resultados:
            if resultado:
                noticias_consolidadas.append(resultado)

    print(f"\n 🎯 Total consolidado do JSON: {len(noticias_consolidadas)} notícias válidas de `{nome_arquivo}`")
    return noticias_consolidadas

def gerar_relatorio_docx(relatorios_finais: List[Dict[str, Any]], pasta_saida: str, stats_funil: Dict[str, int] = None):
    """Gera um documento Word (.docx) interativo com títulos recolhíveis."""
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d_%Hh%Mm")
        nome_arquivo = os.path.join(pasta_saida, f"Relatorio_Consolidado_{timestamp}.docx")
        print(f"\n✍️  Gerando Relatório DOCX: `{nome_arquivo}`")

        # NOVO: Ordenar relatórios por prioridade antes de gerar o DOCX
        ordem_prioridade = {'P1_CRITICO': 1, 'P2_ESTRATEGICO': 2, 'P3_MONITORAMENTO': 3}
        relatorios_finais_ordenados = sorted(relatorios_finais, key=lambda x: ordem_prioridade.get(x.get('prioridade', 'P3_MONITORAMENTO'), 3))

        doc = Document()
        doc.add_heading('Relatório Consolidado de Notícias', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data = doc.add_paragraph()
        p_data.add_run(f'Gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}').italic = True
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar linha do funil de processamento
        if stats_funil:
            p_funil = doc.add_paragraph()
            p_funil.add_run('Funil de Processamento: ').bold = True
            p_funil.add_run(f"{stats_funil.get('noticias_extraidas', 0)} notícias extraídas → "
                           f"{stats_funil.get('grupos_faticos', 0)} eventos consolidados → "
                           f"{stats_funil.get('resumos_finais', 0)} resumos finais")
            p_funil.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Espaçamento

        # NOVO: Adiciona parágrafo instrutivo
        p_instrucao = doc.add_paragraph()
        p_instrucao.add_run('Guia de Leitura: ').bold = True
        p_instrucao.add_run('Clique na seta ao lado de cada título para expandir ou recolher o resumo. '
                          'Use o "Painel de Navegação" (Exibir → Painel de Navegação) para navegar pelo documento.').italic = True
        
        doc.add_page_break()

        doc.add_paragraph()

        # Agrupamento dos resumos por TAG (usando relatórios ordenados)
        TAG_ORDER = [
            'Governo e Politica',
            'Economia e Tecnologia',
            'Judicionario',
            'Empresas Privadas'
        ]
        
        grupos_por_tag: Dict[str, List[Dict[str, Any]]] = {tag: [] for tag in TAG_ORDER}
        for rel in relatorios_finais_ordenados:
            tag_val = rel.get('tag')
            if tag_val in grupos_por_tag:
                grupos_por_tag[tag_val].append(rel)

        doc.add_heading('Índice de Notícias', level=1)
        for tag in TAG_ORDER:
            itens_tag = grupos_por_tag.get(tag, [])
            if not itens_tag:
                continue
            doc.add_paragraph(f'{tag}:', style='Heading 3')
            for relatorio in itens_tag:
                doc.add_paragraph(relatorio["titulo_final"], style='List Number')

        doc.add_page_break()

        doc.add_heading('Análises Consolidadas', level=1)
        for tag in TAG_ORDER:
            itens_tag = grupos_por_tag.get(tag, [])
            if not itens_tag:
                continue
            doc.add_heading(tag, level=2)
            for i, relatorio in enumerate(itens_tag, 1):
                # ALTERAÇÃO PRINCIPAL: Troca 'level=3' por 'style='Heading 3''
                titulo_p = doc.add_paragraph(style='Heading 3')
                titulo_p.add_run(f'{i}. {relatorio["titulo_final"]}')
                
                # O resumo continua como parágrafo normal
                doc.add_paragraph(relatorio['resumo_final'], style='Body Text')

                p_fontes = doc.add_paragraph()
                p_fontes.add_run('Fontes: ').bold = True
                
                fontes_formatadas_unicas = set() # Usar um set para garantir fontes únicas
                
                for fonte in relatorio.get("fontes", []):
                    jornal = fonte.get('jornal', 'N/A').replace('.pdf', '')
                    pagina_str = ""
                    autor_str = ""
                    
                    # Adicionar página apenas se for válida
                    pagina = fonte.get('pagina')
                    if pagina and str(pagina).lower() != 'n/a':
                        pagina_str = f"pág. {pagina}"
                        
                    # Adicionar autor apenas se for válido
                    autor = fonte.get('autor')
                    if autor and str(autor).lower() != 'n/a':
                        autor_str = f"por {autor}"
                        
                    # Montar a string da fonte de forma limpa
                    partes_fonte = [item for item in [pagina_str, autor_str] if item]
                    if partes_fonte:
                        fonte_str = f"{jornal} ({', '.join(partes_fonte)})"
                    else:
                        fonte_str = jornal
                        
                    fontes_formatadas_unicas.add(fonte_str)
                    
                p_fontes.add_run(" | ".join(sorted(list(fontes_formatadas_unicas)))).font.size = Pt(9)
                p_fontes.runs[-1].italic = True
                doc.add_paragraph()

        doc.save(nome_arquivo)
        print(f"✅ Relatório Final Gerado com sucesso em `{nome_arquivo}`")

    except Exception as e:
        print(f"❌ Erro Crítico ao gerar o arquivo DOCX: {e}")


# ==============================================================================
# 7. FUNÇÃO AUXILIAR PARA PROCESSAMENTO PARALELO DE ARQUIVOS
# ==============================================================================

def processar_um_arquivo(args):
    """
    Função worker para processar um único arquivo (PDF ou JSON).
    Projetada para ser chamada em paralelo.
    """
    caminho_arquivo, client, generation_config_decision = args
    nome_arquivo = os.path.basename(caminho_arquivo)
    print(f"\n--- \n📰 Processando arquivo: `{nome_arquivo}`")
    
    # Lógica para JSON (com cache implementado)
    if caminho_arquivo.endswith('.json'):
        nome_cache = f"{limpar_nome_arquivo(nome_arquivo)}_processado.json"
        caminho_cache = os.path.join(CACHE_DIRECTORY, nome_cache)

        if os.path.exists(caminho_cache):
            try:
                print(f"♻️  Tentando carregar do cache: `{caminho_cache}`")
                with open(caminho_cache, 'r', encoding='utf-8') as f:
                    noticias_cache = json.load(f)
                
                # Usar função de migração para validar e migrar dados antigos
                noticias_validadas_cache = validar_e_migrar_cache(noticias_cache, nome_arquivo)
                
                # Se a validação passou sem erros e o cache não está vazio:
                if noticias_validadas_cache:
                    print(f"✅ Cache HIT: {len(noticias_validadas_cache)} notícias carregadas com sucesso.")
                    
                    # Se houve migração, salva o cache atualizado
                    if len(noticias_cache) != len(noticias_validadas_cache) or any('relevance_score' not in n for n in noticias_cache):
                        with open(caminho_cache, 'w', encoding='utf-8') as f:
                            json.dump(noticias_validadas_cache, f, ensure_ascii=False, indent=2)
                        print(f"💾 Cache migrado e atualizado para `{nome_arquivo}`.")
                    
                    return noticias_validadas_cache
                else:
                    # O cache estava vazio ou todos os dados eram inválidos.
                    print("⚠️ Aviso: Cache vazio ou inválido. Reprocessando.")

            except Exception as e:
                print(f"⚠️ Aviso: Falha crítica ao ler cache `{nome_arquivo}`. Reprocessando. Erro: {e}")

        # Se o cache falhou ou não existe, processa o JSON
        noticias_validadas_json = processar_arquivo_json(caminho_arquivo, client, generation_config_decision)
        
        # Salva no cache e retorna
        if noticias_validadas_json:
            with open(caminho_cache, 'w', encoding='utf-8') as f:
                json.dump(noticias_validadas_json, f, ensure_ascii=False, indent=2)
            print(f"💾 Cache SAVE: Resultados de `{nome_arquivo}` salvos.")
            return noticias_validadas_json
        else:
            print(f"✅ Nenhuma notícia relevante encontrada em `{nome_arquivo}`.")
            return []

    # Lógica para PDF (incluindo o novo bloco de cache corrigido)
    elif caminho_arquivo.endswith('.pdf'):
        nome_cache = f"{limpar_nome_arquivo(nome_arquivo)}.json"
        caminho_cache = os.path.join(CACHE_DIRECTORY, nome_cache)

        if os.path.exists(caminho_cache):
            try:
                print(f"♻️  Tentando carregar do cache: `{caminho_cache}`")
                with open(caminho_cache, 'r', encoding='utf-8') as f:
                    noticias_cache = json.load(f)
                
                # Usar função de migração para validar e migrar dados antigos
                noticias_validadas_cache = validar_e_migrar_cache(noticias_cache, nome_arquivo)
                
                # Se a validação passou sem erros e o cache não está vazio:
                if noticias_validadas_cache:
                    print(f"✅ Cache HIT: {len(noticias_validadas_cache)} notícias carregadas com sucesso.")
                    
                    # Se houve migração, salva o cache atualizado
                    if len(noticias_cache) != len(noticias_validadas_cache) or any('relevance_score' not in n for n in noticias_cache):
                        with open(caminho_cache, 'w', encoding='utf-8') as f:
                            json.dump(noticias_validadas_cache, f, ensure_ascii=False, indent=2)
                        print(f"💾 Cache migrado e atualizado para `{nome_arquivo}`.")
                    
                    return noticias_validadas_cache
                else:
                    # O cache estava vazio ou todos os dados eram inválidos.
                    print("⚠️ Aviso: Cache vazio ou inválido. Reprocessando.")

            except Exception as e:
                print(f"⚠️ Aviso: Falha crítica ao ler cache `{nome_arquivo}`. Reprocessando. Erro: {e}")

        # Se o cache falhou ou não existe, processa o PDF
        num_paginas = contar_paginas_pdf(caminho_arquivo)
        usar_chunking = num_paginas > LIMITE_PAGINAS_CHUNKING
        
        noticias_validadas_pdf = []
        if usar_chunking:
            print(f"📊 PDF grande detectado ({num_paginas} páginas) - usando chunking.")
            noticias_validadas_pdf = processar_pdf_com_chunking(caminho_arquivo, client, generation_config_decision)
        else:
            print(f"📄 PDF normal ({num_paginas} páginas) - processamento direto.")
            try:
                uploaded_file = client.files.upload(file=caminho_arquivo)
                while uploaded_file.state.name == "PROCESSING":
                    time.sleep(5)
                    uploaded_file = client.files.get(name=uploaded_file.name)
                
                if uploaded_file.state.name == "ACTIVE":
                    response = client.models.generate_content(
                        model='gemini-2.0-flash-lite',
                        contents=[uploaded_file, PROMPT_EXTRACAO_PERMISSIVO_V8],
                        config=generation_config_decision
                    )
                    response_text = response.text
                    noticias_brutas = extrair_json_da_resposta(response_text)
                    
                    if noticias_brutas and isinstance(noticias_brutas, list):
                        for noticia_data in noticias_brutas:
                            try:
                                noticia_data['jornal'] = nome_arquivo
                                
                                # CORREÇÃO DE TAG ANTES DA VALIDAÇÃO
                                if 'tag' in noticia_data:
                                    noticia_data['tag'] = corrigir_tag_invalida(noticia_data['tag'])
                                
                                noticias_validadas_pdf.append(Noticia(**noticia_data).model_dump())
                            except ValidationError as e:
                                print(f"   ⚠️ Notícia do PDF DESCARTADA: {[error['loc'][0] for error in e.errors()]}")
                    client.files.delete(name=uploaded_file.name)
                else:
                     print(f"❌ Erro no upload do PDF: {uploaded_file.state.name}")
            except Exception as e:
                print(f"❌ ERRO CRÍTICO ao processar PDF `{nome_arquivo}`: {e}")

        # Salva no cache e retorna
        if noticias_validadas_pdf:
            with open(caminho_cache, 'w', encoding='utf-8') as f:
                json.dump(noticias_validadas_pdf, f, ensure_ascii=False, indent=2)
            print(f"💾 Cache SAVE: Resultados de `{nome_arquivo}` salvos.")
            return noticias_validadas_pdf
        else:
            print(f"✅ Nenhuma notícia relevante encontrada em `{nome_arquivo}`.")
            return []
    return []

def gerar_resumo_para_grupo(args):
    """
    Função worker para gerar o resumo final de um único evento com profundidade variável.
    Inclui retry automático e melhor tratamento de erros.
    """
    grupo, i, total, client, generation_config_text = args
    tema = grupo.get('tema_principal', f'Grupo {i}')
    prioridade = grupo.get('prioridade')
    print(f"  -> Resumindo evento {i}/{total} (Prioridade: {prioridade}): {tema}")
    
    # NOVO: Sistema de retry com múltiplas tentativas
    MAX_TENTATIVAS = 3
    
    for tentativa in range(1, MAX_TENTATIVAS + 1):
        try:
            # NOVO: Determinar o nível de detalhe com base na prioridade
            mapa_detalhe = {
                'P1_CRITICO': 'Executivo (P1_CRITICO)',
                'P2_ESTRATEGICO': 'Padrão (P2_ESTRATEGICO)',
                'P3_MONITORAMENTO': 'Conciso (P3_MONITORAMENTO)'
            }
            nivel_de_detalhe = mapa_detalhe.get(prioridade, 'Padrão (P2_ESTRATEGICO)')

            # Passar o 'tema_principal' no JSON para o prompt
            dados_para_resumir = {
                "tema_principal": tema,
                "noticias": grupo.get("noticias", [])
            }
            dados_json_str = json.dumps(dados_para_resumir, ensure_ascii=False, indent=2)

            # --- CORREÇÃO PRINCIPAL AQUI ---
            # Unificar instrução e dados em um único prompt.
            prompt_completo = PROMPT_RESUMO_FINAL_V3.format(
                NIVEL_DE_DETALHE=nivel_de_detalhe,
                DADOS_DO_GRUPO=dados_json_str  # Injeta o JSON como string
            )
            # --- FIM DA CORREÇÃO ---
            
            # NOVO: Configuração mais conservadora para resumos
            config_resumo = types.GenerateContentConfig(
                temperature=0.3,  # Ainda mais baixa para consistência
                top_p=0.9,
                top_k=20,
                max_output_tokens=8192,  # Limite menor para evitar truncamento
            )
            
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                # Envia apenas UMA string com tudo dentro
                contents=[prompt_completo],
                config=config_resumo  # Usa config específica
            )
            
            # NOVO: Verificação de resposta válida antes de extrair JSON
            if not hasattr(response, 'text') or not response.text:
                if tentativa < MAX_TENTATIVAS:
                    print(f"     ⚠️ Tentativa {tentativa}: Resposta vazia da API. Tentando novamente...")
                    time.sleep(1)
                    continue
                else:
                    print(f"     ❌ Falha na tentativa {tentativa}: API retornou resposta vazia.")
                    return None
            
            # NOVO: Tratamento robusto de JSON com debug
            resumo_bruto = extrair_json_da_resposta(response.text)
            
            if not resumo_bruto:
                if tentativa < MAX_TENTATIVAS:
                    print(f"     ⚠️ Tentativa {tentativa}: JSON inválido. Tentando novamente...")
                    time.sleep(1)
                    continue
                else:
                    print(f"     ❌ Falha na tentativa {tentativa}: Não foi possível extrair JSON válido.")
                    print(f"     📋 RESPOSTA DA API (debug): {response.text[:200]}...")
                    return None

            # Validação da estrutura do resumo
            if not isinstance(resumo_bruto, dict) or 'titulo_final' not in resumo_bruto or 'resumo_final' not in resumo_bruto:
                if tentativa < MAX_TENTATIVAS:
                    print(f"     ⚠️ Tentativa {tentativa}: Estrutura JSON incompleta. Tentando novamente...")
                    time.sleep(1)
                    continue
                else:
                    print(f"     ❌ Falha na tentativa {tentativa}: JSON sem campos obrigatórios.")
                    return None

            # Processamento das tags e fontes (já validado)
            tag_counts = {}
            for n in grupo.get('noticias', []):
                tag_val = n.get('tag')
                if tag_val:
                    tag_counts[tag_val] = tag_counts.get(tag_val, 0) + 1
            
            if tag_counts:
                tag_predominante = max(tag_counts, key=tag_counts.get)
                
                fontes_validadas = []
                for noticia in grupo.get('noticias', []):
                    try:
                        fonte = FonteResumo(
                            jornal=noticia.get("jornal"),
                            pagina=noticia.get("pagina"),
                            autor=noticia.get("autor")
                        )
                        fontes_validadas.append(fonte.model_dump())
                    except ValidationError:
                        continue
                
                resumo_bruto["fontes"] = fontes_validadas
                resumo_bruto["tag"] = tag_predominante
                resumo_bruto["prioridade"] = grupo.get("prioridade")
                
                # Validação final com Pydantic
                resumo_obj = ResumoFinal(**resumo_bruto)
                print(f"     ✅ Resumo para '{tema}' gerado e validado com sucesso (tentativa {tentativa}).")
                return resumo_obj.model_dump()
            else:
                print(f"     ❌ Falha: Nenhuma tag válida encontrada no grupo '{tema}'.")
                return None
                
        except json.JSONDecodeError as json_err:
            if tentativa < MAX_TENTATIVAS:
                print(f"     ⚠️ Tentativa {tentativa}: Erro de JSON. Tentando novamente... ({json_err})")
                time.sleep(1)
                continue
            else:
                print(f"     ❌ Falha final na tentativa {tentativa}: Erro JSON persistente para '{tema}': {json_err}")
                return None
                
        except ValidationError as val_err:
            if tentativa < MAX_TENTATIVAS:
                print(f"     ⚠️ Tentativa {tentativa}: Erro de validação. Tentando novamente...")
                time.sleep(1)
                continue
            else:
                print(f"     ❌ Falha final na tentativa {tentativa}: Erro de validação para '{tema}': {val_err}")
                return None
                
        except Exception as e:
            if tentativa < MAX_TENTATIVAS:
                print(f"     ⚠️ Tentativa {tentativa}: Erro geral. Tentando novamente... ({e})")
                time.sleep(1)
                continue
            else:
                print(f"     ❌ Falha final na tentativa {tentativa}: Erro geral para '{tema}': {e}")
                return None
    
    print(f"     ❌ Todas as {MAX_TENTATIVAS} tentativas falharam para o grupo '{tema}'.")
    return None

# ==============================================================================
# 8. FUNÇÃO PRINCIPAL DO PIPELINE (v7.0)
# ==============================================================================

def main(client, generation_config_decision, generation_config_text):
    """
    Executa o pipeline v7.0 com foco em um briefing executivo de duas camadas.
    1. Extrai notícias com score de relevância.
    2. Agrupa por eventos.
    3. RANKING: Separa os eventos, ranqueia os P1 e seleciona o TOP 12.
    4. SUMARIZAÇÃO EM DUAS CAMADAS:
        - TOP 12 Críticos: Resumos de 1 parágrafo.
        - Radar de Monitoramento: Resumos de 1 linha (bullet points).
    5. MONTAGEM DO RELATÓRIO: Gera o DOCX no novo formato.
    """
    print("\n" + "="*80)
    print("🚀 INICIANDO PIPELINE DE BRIEFING EXECUTIVO (v7.0) 🚀")
    print("="*80)

    todas_noticias_extraidas = []
    grupos_de_eventos = []

    # --------------------------------------------------------------------------
    # ETAPA 1: EXTRAÇÃO DE NOTÍCIAS EM PARALELO
    # --------------------------------------------------------------------------
    print("\n" + "-"*22 + " ETAPA 1: Extração em Paralelo " + "-"*23)
    caminhos_pdfs = glob.glob(os.path.join(PDF_DIRECTORY, '*.pdf'))
    caminhos_jsons = glob.glob(os.path.join(PDF_DIRECTORY, '*.json'))
    todos_os_arquivos = caminhos_pdfs + caminhos_jsons

    if not todos_os_arquivos:
        print(f"❌ ERRO: Nenhum arquivo PDF ou JSON encontrado na pasta `{PDF_DIRECTORY}`. Encerrando.")
        return

    MAX_WORKERS_EXTRACAO = 10
    print(f"🚀 Iniciando extração paralela de {len(todos_os_arquivos)} arquivos com até {MAX_WORKERS_EXTRACAO} workers...")
    print(f"📁 Distribuição: {len(caminhos_pdfs)} PDFs e {len(caminhos_jsons)} JSONs")

    tarefas_extr = [(path, client, generation_config_decision) for path in todos_os_arquivos]

    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS_EXTRACAO) as executor:
        resultados_por_arquivo = executor.map(processar_um_arquivo, tarefas_extr)
        
        for lista_de_noticias in resultados_por_arquivo:
            if lista_de_noticias:
                todas_noticias_extraidas.extend(lista_de_noticias)

    print(f"\n🎯 EXTRAÇÃO PARALELA CONCLUÍDA: {len(todas_noticias_extraidas)} notícias totais de {len(todos_os_arquivos)} arquivos.")

    # --------------------------------------------------------------------------
    # ETAPA 2: AGRUPAMENTO CONSOLIDADO DE EVENTOS
    # --------------------------------------------------------------------------
    print("\n" + "-"*21 + " ETAPA 2: Agrupamento Consolidado " + "-"*22)
    if todas_noticias_extraidas:
        noticias_para_agrupar_com_id = [
            {
                "id": i,
                "titulo": n.get("titulo", ""),
                "jornal": n.get("jornal", ""),
                "trecho": (n.get("texto_completo", "")[:300] + "...") if len(n.get("texto_completo", "")) > 300 else n.get("texto_completo", "")
            }
            for i, n in enumerate(todas_noticias_extraidas)
        ]
        mapa_id_para_noticia = {i: n for i, n in enumerate(todas_noticias_extraidas)}

        print(f"🔄 Consolidando {len(noticias_para_agrupar_com_id)} notícias em eventos únicos (em lotes)...")

        # Novo: Agrupamento em lotes com retry e fallback
        tamanho_lote = 120
        max_tentativas = 3
        grupos_por_lote: list[dict] = []

        for inicio in range(0, len(noticias_para_agrupar_com_id), tamanho_lote):
            fim = min(inicio + tamanho_lote, len(noticias_para_agrupar_com_id))
            lote = noticias_para_agrupar_com_id[inicio:fim]
            print(f"   📦 Lote {inicio//tamanho_lote + 1}: itens {inicio}–{fim-1} ({len(lote)} notícias)")

            grupos_brutos = None
            for tentativa in range(1, max_tentativas + 1):
                try:
                    prompt_completo = PROMPT_AGRUPAMENTO_CONSOLIDADO_V2 + "\n\nNOTÍCIAS PARA AGRUPAR:\n" + json.dumps(lote, indent=2, ensure_ascii=False)
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=[prompt_completo],
                        config=generation_config_text
                    )

                    if not hasattr(response, 'text') or not isinstance(response.text, str) or not response.text.strip():
                        print(f"     ⚠️ Tentativa {tentativa}: Resposta vazia do LLM para o lote. Retry...")
                        time.sleep(1)
                        continue

                    grupos_brutos = extrair_json_da_resposta(response.text)
                    if not grupos_brutos or not isinstance(grupos_brutos, list):
                        print(f"     ⚠️ Tentativa {tentativa}: JSON inválido no lote. Retry...")
                        time.sleep(1)
                        grupos_brutos = None
                        continue

                    # Sucesso
                    break
                except Exception as e:
                    print(f"     ⚠️ Tentativa {tentativa}: Erro ao agrupar lote: {e}")
                    time.sleep(1)

            if not grupos_brutos:
                # Fallback: cada item do lote vira seu próprio grupo
                print("     🔁 Fallback ativado: criando grupos unitários para o lote.")
                for item in lote:
                    grupos_por_lote.append({
                        "tema_principal": item.get("titulo") or "Evento",
                        "ids_originais": [item.get("id")]
                    })
            else:
                grupos_por_lote.extend(grupos_brutos)

        # Consolidação dos grupos de todos os lotes em 'grupos_de_eventos'
        # Mescla por 'tema_principal' idêntico para reduzir duplicatas simples
        grupos_por_tema: dict[str, set] = {}
        for grupo in grupos_por_lote:
            tema = (grupo.get("tema_principal") or "Evento").strip()
            ids = grupo.get("ids_originais", []) or []
            if tema not in grupos_por_tema:
                grupos_por_tema[tema] = set()
            grupos_por_tema[tema].update(ids)

        for tema, ids_set in grupos_por_tema.items():
            noticias_consolidadas_grupo = []
            for id_original in ids_set:
                noticia_original = mapa_id_para_noticia.get(id_original)
                if noticia_original:
                    noticias_consolidadas_grupo.append(noticia_original)
            if noticias_consolidadas_grupo:
                grupos_de_eventos.append({
                    "tema_principal": tema,
                    "noticias": noticias_consolidadas_grupo
                })

        print(f"✅ {len(grupos_de_eventos)} eventos únicos criados após consolidação em lotes.")
    else:
        print("⚠️ Aviso: Nenhuma notícia extraída, pulando agrupamentos.")

    # --------------------------------------------------------------------------
    # ETAPA 3: SEPARAÇÃO E RANKING DOS GRUPOS
    # --------------------------------------------------------------------------
    print("\n" + "-"*23 + " ETAPA 3: Ranking e Seleção de Grupos " + "-"*24)
    
    grupos_p1_candidatos = []
    grupos_monitoramento = []

    if not grupos_de_eventos:
        print("❌ Nenhum grupo de eventos foi criado. Encerrando.")
        return

    for grupo in grupos_de_eventos:
        prioridades = [n.get('prioridade') for n in grupo['noticias']]
        # Calcula o score médio do grupo
        scores = [n.get('relevance_score', 0) for n in grupo['noticias']]
        grupo['score_medio'] = sum(scores) / len(scores) if scores else 0
        
        # Atribui a maior prioridade ao grupo
        if 'P1_CRITICO' in prioridades:
            grupo['prioridade'] = 'P1_CRITICO'
            grupos_p1_candidatos.append(grupo)
        else:
            grupo['prioridade'] = 'P2_ESTRATEGICO' if 'P2_ESTRATEGICO' in prioridades else 'P3_MONITORAMENTO'
            grupos_monitoramento.append(grupo)

    # Ranqueia os grupos P1 pelo score e seleciona o TOP 12
    grupos_p1_candidatos.sort(key=lambda g: g['score_medio'], reverse=True)
    top_12_criticos = grupos_p1_candidatos[:12]
    
    # O resto dos P1 vai para o monitoramento também
    grupos_monitoramento.extend(grupos_p1_candidatos[12:])
    
    print(f"🎯 Seleção Crítica: {len(top_12_criticos)} eventos P1 selecionados para análise detalhada.")
    print(f"📡 Radar de Monitoramento: {len(grupos_monitoramento)} eventos para resumos de 1 linha.")
    
    # --------------------------------------------------------------------------
    # ETAPA 4: GERAÇÃO DE RESUMOS (DUAS CAMADAS)
    # --------------------------------------------------------------------------
    print("\n" + "-"*21 + " ETAPA 4: Geração de Resumos (2 Camadas) " + "-"*21)
    
    # Camada 1: Resumos detalhados para o TOP 12
    relatorios_criticos = []
    if top_12_criticos:
        print(f"📝 Gerando resumos detalhados para {len(top_12_criticos)} eventos críticos...")
        
        # Gerar resumos críticos em paralelo
        MAX_WORKERS_RESUMO = 15
        tarefas_criticos = [(grupo, client, generation_config_text) for grupo in top_12_criticos]
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS_RESUMO) as executor:
            resultados_criticos = executor.map(gerar_resumo_critico, tarefas_criticos)
            relatorios_criticos = [res for res in resultados_criticos if res is not None]

    # Camada 2: Bullet points para o Radar
    relatorios_radar = {}  # Dicionário para agrupar por seção
    if grupos_monitoramento:
        print(f"⚡ Gerando bullet points para {len(grupos_monitoramento)} eventos do radar...")
        
        # Gerar bullet points em paralelo
        tarefas_radar = [(grupo, client, generation_config_text) for grupo in grupos_monitoramento]
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS_RESUMO) as executor:
            resultados_radar = executor.map(gerar_resumo_radar, tarefas_radar)
            
            # Agrupar por seções
            relatorios_radar = {
                'Empresas Privadas': [],
                'Economia e Politica': [],
                'Legislativo e Judiciario': [],
                'Internacional': [],
                'Tecnologia': []
            }
            
            for resultado in resultados_radar:
                if resultado and 'bullet_point' in resultado and 'tag' in resultado:
                    tag = resultado['tag']
                    bullet = resultado['bullet_point']
                    
                    # Mapear tags para seções
                    if tag == 'Empresas Privadas':
                        relatorios_radar['Empresas Privadas'].append(bullet)
                    elif tag == 'Governo e Politica':
                        relatorios_radar['Economia e Politica'].append(bullet)
                    elif tag == 'Judicionario':
                        relatorios_radar['Legislativo e Judiciario'].append(bullet)
                    elif tag == 'Economia e Tecnologia':
                        # Dividir entre tecnologia e internacional baseado no conteúdo
                        if any(tech_word in bullet.lower() for tech_word in ['ia', 'inteligência artificial', 'tecnologia', 'software', 'app', 'digital', 'crypto', 'bitcoin']):
                            relatorios_radar['Tecnologia'].append(bullet)
                        else:
                            relatorios_radar['Internacional'].append(bullet)

    # --------------------------------------------------------------------------
    # ETAPA 5: GERAÇÃO DO RELATÓRIO FINAL EM DOCX
    # --------------------------------------------------------------------------
    print("\n" + "-"*25 + " ETAPA 5: Geração do Relatório " + "-"*26)
    
    # Coletar nomes dos jornais analisados (agora o campo jornal já tem a fonte correta)
    nomes_jornais = set()
    for n in todas_noticias_extraidas:
        jornal_name = n.get('jornal', '')
        # Remove extensões de arquivo se existirem
        jornal_limpo = jornal_name.replace('.pdf', '').replace('.json', '')
        if jornal_limpo and jornal_limpo != 'N/A':
            nomes_jornais.add(jornal_limpo)

    stats_funil = {
        'jornais_analisados': ", ".join(sorted(list(nomes_jornais))),
        'noticias_analisadas': len(todas_noticias_extraidas),
        'eventos_agrupados': len(grupos_de_eventos),
        'resumos_criticos': len(relatorios_criticos)
    }

    gerar_relatorio_docx_v2(
        relatorios_criticos,
        relatorios_radar,
        OUTPUT_DIRECTORY,
        stats_funil
    )
    
    print("\n✅ PIPELINE DE BRIEFING EXECUTIVO v7.0 CONCLUÍDO.\n")

# ==============================================================================
# 9. FUNÇÕES AUXILIARES PARA O NOVO PIPELINE v7.0
# ==============================================================================

def gerar_resumo_critico(tarefa):
    """
    Gera um resumo crítico de 1 parágrafo para eventos do TOP 12.
    
    Args:
        tarefa: Tupla (grupo, client, generation_config_text)
    
    Returns:
        Dict com resumo crítico ou None em caso de erro
    """
    grupo, client, generation_config_text = tarefa
    
    try:
        # Preparar dados do grupo
        dados_grupo = {
            "tema_principal": grupo.get("tema_principal", ""),
            "noticias": grupo.get("noticias", [])
        }
        
        prompt_completo = PROMPT_RESUMO_CRITICO_V1.format(
            DADOS_DO_GRUPO=json.dumps(dados_grupo, indent=2, ensure_ascii=False)
        )
        
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt_completo],
            config=generation_config_text
        )
        
        resultado = extrair_json_da_resposta(response.text)
        
        if resultado and 'resumo_final' in resultado:
            # Preparar fontes
            fontes = []
            for noticia in grupo.get("noticias", []):
                # --- ALTERAÇÃO AQUI ---
                # Adicione a URL ao dicionário da fonte
                fontes.append({
                    'jornal': noticia.get('jornal', 'N/A'),
                    'pagina': noticia.get('pagina', 'N/A'),
                    'autor': noticia.get('autor', 'N/A'),
                    'url': noticia.get('url') # Adiciona a URL capturada
                })
            
            return {
                'titulo_final': grupo.get("tema_principal", ""),
                'resumo_final': resultado['resumo_final'],
                'fontes': fontes,
                'tag': grupo.get("noticias", [{}])[0].get('tag', 'Empresas Privadas'),
                'prioridade': 'P1_CRITICO'
            }
        
    except Exception as e:
        print(f"❌ Erro ao gerar resumo crítico para '{grupo.get('tema_principal', '')}': {e}")
    
    return None

def gerar_resumo_radar(tarefa):
    """
    Gera um bullet point de 1 linha para eventos do radar de monitoramento.
    
    Args:
        tarefa: Tupla (grupo, client, generation_config_text)
    
    Returns:
        Dict com bullet point e tag ou None em caso de erro
    """
    grupo, client, generation_config_text = tarefa
    
    try:
        # Preparar dados do grupo
        dados_grupo = {
            "tema_principal": grupo.get("tema_principal", ""),
            "noticias": grupo.get("noticias", [])
        }
        
        prompt_completo = PROMPT_RADAR_MONITORAMENTO_V1.format(
            DADOS_DO_GRUPO=json.dumps(dados_grupo, indent=2, ensure_ascii=False)
        )
        
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt_completo],
            config=generation_config_text
        )
        
        resultado = extrair_json_da_resposta(response.text)
        
        if resultado and 'bullet_point' in resultado:
            return {
                'bullet_point': resultado['bullet_point'],
                'tag': grupo.get("noticias", [{}])[0].get('tag', 'Empresas Privadas')
            }
        
    except Exception as e:
        print(f"❌ Erro ao gerar bullet point para '{grupo.get('tema_principal', '')}': {e}")
    
    return None

def gerar_relatorio_docx_v2(
    relatorios_criticos: List[Dict[str, Any]],
    relatorios_radar: Dict[str, List[str]],
    pasta_saida: str,
    stats_funil: Dict[str, any]
):
    """
    Gera o Briefing Executivo no formato de duas camadas, com formatação detalhada das fontes.
    1. Análises Críticas (TOP 12) com resumos e fontes detalhadas.
    2. Radar de Monitoramento com bullet points.
    """
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d_%Hh%Mm")
        nome_arquivo = os.path.join(pasta_saida, f"Briefing_Executivo_{timestamp}.docx")
        print(f"\n✍️ Montando Briefing Executivo com fontes detalhadas: `{nome_arquivo}`")
        doc = Document()
        
        # --- CABEÇALHO (sem alterações) ---
        doc.add_heading('Relatório Consolidado de Notícias', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data = doc.add_paragraph()
        p_data.add_run(f'Gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}').italic = True
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        p_funil = doc.add_paragraph()
        p_funil.add_run('Jornais analisados: ').bold = True
        p_funil.add_run(f"{stats_funil.get('jornais_analisados', 'N/A')}\n")
        p_funil.add_run('Notícias Analisadas: ').bold = True
        p_funil.add_run(f"{stats_funil.get('noticias_analisadas', 0)} → ")
        p_funil.add_run('Agrupadas em: ').bold = True
        p_funil.add_run(f"{stats_funil.get('eventos_agrupados', 0)} → ")
        p_funil.add_run('Resumidas: ').bold = True
        p_funil.add_run(str(stats_funil.get('resumos_criticos', 0)))
        p_funil.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- CAMADA 1: ANÁLISES CRÍTICAS (TOP 12) ---
        if relatorios_criticos:
            doc.add_heading('Principais Noticias', level=1)
            for i, relatorio in enumerate(relatorios_criticos, 1):
                doc.add_heading(f"{i}. {relatorio['titulo_final']}", level=2)
                p_resumo = doc.add_paragraph(relatorio.get('resumo_final', 'N/A'))
                p_resumo.paragraph_format.space_after = Pt(12)
                
                # ======================================================================
                # INÍCIO DA CORREÇÃO DEFINITIVA NO TRATAMENTO DAS FONTES
                # ======================================================================
                p_fontes = doc.add_paragraph()
                p_fontes.add_run('Fontes: ').bold = True
                
                # Não vamos mais usar um 'set' para não perder detalhes.
                # Processamos cada fonte individualmente da lista.
                lista_de_fontes = relatorio.get('fontes', [])
                for idx, fonte in enumerate(lista_de_fontes):
                    # Limpa o nome do jornal para exibição
                    jornal_limpo = str(fonte.get('jornal', 'N/A')).replace('.pdf','').replace('.json','')
                    url = fonte.get('url')
                    
                    # Lógica para fontes de SITES/JSON (que possuem URL)
                    if url:
                        # Cria um hiperlink com o nome do jornal como texto clicável
                        add_hyperlink(p_fontes, jornal_limpo, url)
                    
                    # Lógica para fontes de PDFS (que NÃO possuem URL)
                    else:
                        pagina = fonte.get('pagina')
                        autor = fonte.get('autor')
                        
                        # Constrói a string de detalhes (página, autor) de forma dinâmica
                        partes_detalhe = []
                        if pagina and str(pagina).strip().lower() not in ['n/a', '']:
                            partes_detalhe.append(f"pág. {str(pagina).strip()}")
                        if autor and str(autor).strip().lower() not in ['n/a', '']:
                            partes_detalhe.append(f"por {str(autor).strip()}")
                        
                        # Formata a string final
                        if partes_detalhe:
                            detalhes_str = ", ".join(partes_detalhe)
                            texto_fonte = f"{jornal_limpo.upper()} ({detalhes_str})"
                        else:
                            texto_fonte = jornal_limpo.upper()
                        
                        p_fontes.add_run(texto_fonte).italic = True
                    
                    # Adiciona um separador entre as fontes, mas não após a última
                    if idx < len(lista_de_fontes) - 1:
                        p_fontes.add_run(' | ').italic = True
                
                doc.add_paragraph() # Adiciona um espaço após a linha de fontes
                # ======================================================================
                # FIM DA CORREÇÃO DEFINITIVA
                # ======================================================================
        
        # --- CAMADA 2: RADAR DE MONITORAMENTO (sem alterações) ---
        if relatorios_radar:
            doc.add_heading('Radar de Monitoramento', level=1)
            SECTIONS_ORDER = ['Empresas Privadas', 'Economia e Politica', 'Legislativo e Judiciario', 'Internacional', 'Tecnologia']
            for section in SECTIONS_ORDER:
                if section in relatorios_radar and relatorios_radar[section]:
                    doc.add_heading(section, level=2)
                    for bullet_point in relatorios_radar[section][:12]:
                        if bullet_point.startswith('•'):
                            bullet_point = bullet_point[1:].strip()
                        if ':' in bullet_point:
                            partes = bullet_point.split(':', 1)
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(partes[0] + ':').bold = True
                            p.add_run(partes[1])
                        else:
                            p = doc.add_paragraph(bullet_point, style='List Bullet')
                        p.paragraph_format.space_after = Pt(6)
                    doc.add_paragraph()

        doc.save(nome_arquivo)
        print(f"✅ Briefing Executivo Gerado com sucesso em `{nome_arquivo}`")

    except Exception as e:
        print(f"❌ Erro Crítico ao gerar o arquivo DOCX: {e}")

# ==============================================================================
# 10. PONTO DE ENTRADA DO SCRIPT
# ==============================================================================

if __name__ == "__main__":
    # Verificação de dependências
    if not verificar_dependencias():
        exit(1)
    
    for diretorio in [PDF_DIRECTORY, OUTPUT_DIRECTORY, CACHE_DIRECTORY]:
        os.makedirs(diretorio, exist_ok=True)

    # Configuração da API Key
    api_key = "AIzaSyAB7hZ1C9t_Tb-q-sXLzLvhqbJUmB3noUE"
    print("🔑 Usando chave de API definida diretamente no código.")
    
    try:
        # Criar cliente com a nova API
        client = genai.Client(api_key=api_key)
        
        # Configuração de IA otimizada para precisão em decisões
        generation_config_decision = types.GenerateContentConfig(
            temperature=0.1,  # Baixa temperatura para decisões precisas
            top_p=0.95,
            top_k=40,
            max_output_tokens=65536,
        )

        # Configuração para tarefas de processamento de texto (agrupar, resumir)
        generation_config_text = types.GenerateContentConfig(
            temperature=0.5,  # Temperatura moderada para criatividade controlada
            top_p=0.95,
            top_k=40,
            max_output_tokens=65536,
        )

        print("⚡️ Inicializando cliente de IA com configurações específicas...")

        model_name1 = 'gemini-2.5-flash-lite'
        model_name2 = 'gemini-2.5-flash'


        print(f"   - Modelo para Decisões (Extrair/Filtrar): {model_name1} (temp={generation_config_decision.temperature})")
        print(f"   - Modelo para Texto (Agrupar/Resumir): {model_name2} (temp={generation_config_text.temperature})")

        # Passa o cliente e as configurações para a função main
        main(client=client, generation_config_decision=generation_config_decision, generation_config_text=generation_config_text)

    except Exception as e:
        print(f"\n❌ Ocorreu um erro fatal durante a inicialização ou execução: {e}")
        print("   Verifique sua chave de API, conexão com a internet e permissões de pasta.")
        exit(1) 